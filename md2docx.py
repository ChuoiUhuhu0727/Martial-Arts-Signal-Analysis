"""Convert the Subsystem A report markdown into a formatted .docx."""
import base64
import io
import os
import re
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

BODY_FONT = "Calibri"
CODE_FONT = "Consolas"
ACCENT = RGBColor(0x1F, 0x3B, 0x63)
MUTED = RGBColor(0x55, 0x5F, 0x6D)
CODE_BG = "F2F4F7"
HEAD_BG = "E8EDF4"
QUOTE_BG = "FBF6E8"


def shade(el, fill):
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), fill)
    el.append(sh)


def border(par, edge="bottom", sz=6, color="B8C4D6"):
    pPr = par._p.get_or_add_pPr()
    bdr = pPr.find(qn("w:pBdr"))
    if bdr is None:
        bdr = OxmlElement("w:pBdr")
        pPr.append(bdr)
    e = OxmlElement(f"w:{edge}")
    e.set(qn("w:val"), "single")
    e.set(qn("w:sz"), str(sz))
    e.set(qn("w:space"), "4")
    e.set(qn("w:color"), color)
    bdr.append(e)


INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\*[^*\n]+?\*)")


def add_runs(par, text, bold=False, italic=False, size=None, color=None):
    """Emit `text` as formatted runs.

    Markers nest in real markdown -- **`code` inside bold** is common -- so stripping
    one layer and emitting immediately would print the inner markers literally. Recurse
    on the stripped content instead, carrying the accumulated bold/italic state, and
    only emit a run once no marker is left to peel.
    """
    for tok in INLINE.split(text):
        if not tok:
            continue

        if tok.startswith("**") and tok.endswith("**") and len(tok) > 4:
            add_runs(par, tok[2:-2], bold=True, italic=italic, size=size, color=color)
            continue
        if tok.startswith("*") and tok.endswith("*") and len(tok) > 2:
            add_runs(par, tok[1:-1], bold=bold, italic=True, size=size, color=color)
            continue

        code = tok.startswith("`") and tok.endswith("`") and len(tok) > 2
        if code:
            tok = tok[1:-1]

        r = par.add_run(tok)
        r.bold, r.italic = bold, italic
        r.font.name = CODE_FONT if code else BODY_FONT
        r.font.size = Pt(9.5) if code else (size or Pt(10.5))
        if code:
            r.font.color.rgb = RGBColor(0xA3, 0x1D, 0x3F)
        elif color:
            r.font.color.rgb = color


def para(doc, space_after=6, space_before=0, indent=0):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = 1.15
    if indent:
        pf.left_indent = Inches(indent)
    return p


def heading(doc, text, level):
    sizes = {1: 19, 2: 14.5, 3: 11.5}
    p = para(doc, space_after=6 if level > 1 else 10,
             space_before=0 if level == 1 else (16 if level == 2 else 11))
    add_runs(p, text, bold=True, size=Pt(sizes[level]), color=ACCENT)
    p.paragraph_format.keep_with_next = True
    if level <= 2:
        border(p)
    return p


def code_block(doc, lines):
    p = para(doc, space_after=2, space_before=6, indent=0.12)
    p.paragraph_format.line_spacing = 1.0
    # Keep a block intact rather than letting a diagram or matrix split across pages.
    p.paragraph_format.keep_together = True
    shade(p._p.get_or_add_pPr(), CODE_BG)
    for n, ln in enumerate(lines):
        r = p.add_run(ln)
        r.font.name = CODE_FONT
        r.font.size = Pt(8.5)
        if n < len(lines) - 1:
            r.add_break()


def quote_block(doc, lines):
    for n, ln in enumerate(lines):
        p = para(doc, space_after=2 if n < len(lines) - 1 else 8,
                 space_before=6 if n == 0 else 0, indent=0.25)
        shade(p._p.get_or_add_pPr(), QUOTE_BG)
        add_runs(p, ln, size=Pt(10), color=MUTED)


def split_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def add_table(doc, rows, aligns):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    for ri, row in enumerate(rows):
        for ci, cell in enumerate(row):
            c = t.cell(ri, ci)
            c.text = ""
            p = c.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            if aligns[ci] == "right":
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            add_runs(p, cell, bold=(ri == 0), size=Pt(9.5))
            if ri == 0:
                shade(c._tc.get_or_add_tcPr(), HEAD_BG)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def bullet(doc, text, ordered=False, level=0):
    style = "List Number" if ordered else "List Bullet"
    p = doc.add_paragraph(style=style)
    pf = p.paragraph_format
    pf.space_after = Pt(3)
    pf.line_spacing = 1.15
    pf.left_indent = Inches(0.3 + 0.25 * level)
    add_runs(p, text)
    return p


def page_number_footer(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.font.name = BODY_FONT
    r.font.size = Pt(9)
    r.font.color.rgb = MUTED
    for instr in ("begin", "PAGE", "end"):
        el = OxmlElement("w:fldChar") if instr != "PAGE" else OxmlElement("w:instrText")
        if instr == "PAGE":
            el.set(qn("xml:space"), "preserve")
            el.text = " PAGE "
        else:
            el.set(qn("w:fldCharType"), instr)
        r._r.append(el)


def main(md_path, out_path):
    raw = open(md_path, encoding="utf-8").read().splitlines()

    image_bytes = None
    body = []
    for ln in raw:
        m = re.match(r"^\[image1\]:\s*<data:image/png;base64,(.+)>\s*$", ln)
        if m:
            image_bytes = base64.b64decode(m.group(1))
        else:
            body.append(ln)

    doc = Document()
    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Inches(1.0)
    sec.top_margin = sec.bottom_margin = Inches(0.9)
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10.5)
    page_number_footer(sec)

    i = 0
    while i < len(body):
        line = body[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("```"):
            j = i + 1
            buf = []
            while j < len(body) and not body[j].strip().startswith("```"):
                buf.append(body[j])
                j += 1
            code_block(doc, buf)
            i = j + 1
            continue

        if stripped == "---":
            p = para(doc, space_after=2, space_before=2)
            border(p, color="D6DEE9")
            i += 1
            continue

        if stripped.startswith("!["):
            # Two forms: a reference-style image backed by an inline base64 blob, and
            # ![caption](path) pointing at a file next to the markdown source.
            m_img = re.match(r"^!\[(.*?)\]\((.+?)\)\s*$", stripped)
            if m_img:
                caption, rel = m_img.group(1), m_img.group(2)
                src = os.path.normpath(os.path.join(os.path.dirname(md_path), rel))
                if os.path.exists(src):
                    p = para(doc, space_after=2, space_before=10)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.keep_with_next = True
                    p.add_run().add_picture(src, width=Inches(6.3))
                    if caption:
                        cp = para(doc, space_after=12, space_before=0)
                        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        add_runs(cp, caption, italic=True, size=Pt(9), color=MUTED)
                else:
                    print(f"  [thiếu ảnh] {src}")
            elif image_bytes:
                p = para(doc, space_after=8, space_before=8)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(io.BytesIO(image_bytes), width=Inches(4.9))
            i += 1
            continue

        if stripped.startswith(">"):
            buf = []
            while i < len(body) and body[i].strip().startswith(">"):
                buf.append(body[i].strip().lstrip(">").strip())
                i += 1
            merged, cur = [], ""
            for ln in buf:
                if not ln:
                    if cur:
                        merged.append(cur)
                        cur = ""
                else:
                    cur = f"{cur} {ln}".strip()
            if cur:
                merged.append(cur)
            quote_block(doc, merged)
            continue

        if stripped.startswith("|") and i + 1 < len(body) and re.match(
                r"^\|[\s:|-]+\|$", body[i + 1].strip()):
            header = split_row(body[i])
            aligns = ["right" if c.strip().endswith(":") and not c.strip().startswith(":")
                      else "left" for c in split_row(body[i + 1])]
            rows = [header]
            j = i + 2
            while j < len(body) and body[j].strip().startswith("|"):
                rows.append(split_row(body[j]))
                j += 1
            add_table(doc, rows, aligns)
            i = j
            continue

        m = re.match(r"^(#{1,3})\s+(.*)$", stripped)
        if m:
            heading(doc, m.group(2).strip(), len(m.group(1)))
            i += 1
            continue

        m = re.match(r"^(\s*)[-*]\s+\[( |x)\]\s+(.*)$", line)
        if m:
            mark = "☒" if m.group(2) == "x" else "☐"
            text, j = m.group(3), i + 1
            while j < len(body) and body[j].startswith("      ") and body[j].strip():
                text += " " + body[j].strip()
                j += 1
            p = para(doc, space_after=3, indent=0.3)
            r = p.add_run(mark + "  ")
            r.font.size = Pt(11)
            add_runs(p, text)
            i = j
            continue

        m = re.match(r"^(\s*)([-*]|\d+\.)\s+(.*)$", line)
        if m:
            lead, marker, text = m.group(1), m.group(2), m.group(3)
            j = i + 1
            cont = len(lead) + len(marker) + 1
            while (j < len(body) and body[j].strip()
                   and body[j].startswith(" " * max(cont - 1, 2))
                   and not re.match(r"^\s*([-*]|\d+\.)\s+", body[j])):
                text += " " + body[j].strip()
                j += 1
            bullet(doc, text, ordered=marker[0].isdigit(), level=len(lead) // 2)
            i = j
            continue

        buf = [stripped]
        j = i + 1
        while j < len(body) and body[j].strip() and not re.match(
                r"^\s*(#{1,3}\s|[-*]\s|\d+\.\s|\||>|```|---$)", body[j]):
            buf.append(body[j].strip())
            j += 1
        add_runs(para(doc), " ".join(buf))
        i = j

    doc.save(out_path)
    print(f"wrote {out_path}")


if __name__ == "__main__":
    main(sys.argv[1], sys.argv[2])
